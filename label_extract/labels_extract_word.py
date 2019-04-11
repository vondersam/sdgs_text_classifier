#https://textract.readthedocs.io/en/stable/
from docx import Document as Docxument
import docx
import subprocess
import re
import os
import json
from string import punctuation
import PyPDF2
from bs4 import BeautifulSoup

MAPPINGS = {
        'g': 'sdg|goals|goal',
        't': 'target',
        'i': 'indicator'
    }


class Text:
    def __init__(self, t):
        self.text = t


class Document:
    """
    Doc containing all paragraphs from .doc, .docx, .pdf
    """
    def __init__(self, filepath):
        self.paragraphs = []
        self.extract_paragraphs(filepath)


    def extract_paragraphs(self, filepath):
        base = os.path.basename(filepath)
        extension = os.path.splitext(base)[1]
        if extension == '.doc':
            self.from_word(filepath)
        if extension == '.pdf':
            self.from_pdf(filepath)
        if extension == '.html':
            self.from_html(filepath)
        else:
            print("Cannot process files in that format")

    def from_word(self, file):
        try:
            # Docx
            self.paragraphs = Docxument(file).paragraphs
        except:
            try:
                # Doc
                extracted_string = subprocess.check_output(['antiword', '-t', file])
                # Decode and split by paragraphs
                extracted_list = extracted_string.decode('utf-8').split('\n\n')
                for paragraph in extracted_list:
                    # Eliminate extra spaces, replace space characters and eliminate break lines
                    processed_paragraph = re.sub(' +', ' ', paragraph.strip().replace('\xa0', ' ').replace('\n', ' '))
                    self.paragraphs.append(Text(processed_paragraph))
            except Exception as e:
                print(e)

    def from_pdf(self, file):
        try:
            with open(file, 'rb') as fi:
                pdfReader = PyPDF2.PdfFileReader(fi)
                for i in range(pdfReader.numPages):
                    extracted_string = pdfReader.getPage(i).extractText()
                    extracted_list = extracted_string.split('\n\n')
                    for line in extracted_list:
                        # Eliminate extra spaces, replace space characters and eliminate break lines
                        processed_text = re.sub(' +', ' ', line.strip().replace('\xa0', ' ').replace('\n', ' '))
                        self.paragraphs.append(Text(processed_text))
        except Exception as e:
            print(e)

    def from_html(self, file):
        with open(file, 'rb') as f:
            soup = BeautifulSoup(f, 'html.parser')
            for paragraph in soup.stripped_strings:
                self.paragraphs.append(Text(paragraph))


def extract_type(type_):
    """
    Extract the type of label
    :param type_:
    :return:
    """
    for key, pattern in MAPPINGS.items():
        if type_.lower() in pattern:
            return key


def extract_labels(type_, numbers):
    """

    :param type_:
    :param numbers:
    :return: list of labels
    """
    labels = []
    label_type = extract_type(type_)
    label_numbers = set([i.strip(punctuation) for i in numbers.split() if 'and' not in i])
    for number in label_numbers:
        labels.append(f'{label_type}_{number}')
    return labels



def trans_labels(labels):
    results = dict()
    for i in range(1,18):
        label = f'g_{i}'
        if label in labels:
            results[label] = True
        else:
            results[label] = False
    return results


def extract(filepath, folders=[], spacy_format=False, save=False):
    documents = []
    for folder in os.listdir(filepath):
        if folder in folders:
            documents += [os.path.join(folder, f) for f in os.listdir(os.path.join(filepath, folder))]
    doc_tracker = {}
    training_set = {}

    for document in documents:
        path = os.path.join(filepath + document)
        print(path)
        doc = Document(path)

        for paragraph in doc.paragraphs:
            text = paragraph.text
            # To avoid extracting Millennium Goals
            if 'millennium' not in text.lower():
                patterns = [
                    f"({MAPPINGS['g']})\s+([,*\s*\d+]+[and]*[\s+\d+]*)",
                    f"({MAPPINGS['t']})(\s+\d+\.[\d+a-d])",
                    f"({MAPPINGS['i']})(\s+\d+\.[\d+a-d]\.\d+)"
                ]
                for pattern in patterns:
                    goals = re.findall(pattern, text, re.I)
                    for type_, numbers in goals:
                        if numbers:
                            labels = extract_labels(type_, numbers)
                            if labels:
                                if text not in training_set:
                                    training_set[text] = {
                                        'cats': labels
                                        #'labels': [],
                                        #'doc_id': document
                                    }
                                training_set[text]['cats'].extend(labels)
                                training_set[text]['cats'] = list(set(training_set[text]['cats']))

                                #if document not in doc_tracker:
                                #    doc_tracker[document] = []
                                #doc_tracker[document].extend(labels)
    if spacy_format:
        """
        final_training = {}
        for k, v in training_set.items():
            final_training[k] = {'cats': trans_labels(v['cats'])}
        """

        # Track documents with one
        # for doc, labels in doc_tracker.items():
        #    if len(set(labels)) == 1:
        #        print(doc, labels)
    if save:
        with open('html.json', 'w') as fo:
            json.dump(training_set, fo)
        print(len(training_set))


main_dir = '/Users/samuelrodriguezmedina/Documents/ir4sdgs/crawl_sdgs/'
folders = ['word', 'other_html', 'pdf']
extract(main_dir, folders, save=True)

'''
- we need to do more processing:
    - delete numbered bullet points starting a paragraph
    - delete noisy symbols such as bar |
    - what should we do with the numbers, label them? <number>
- if one label happens more than once, should they be weighted? Or maybe just use a set, do not repeat labels
- if only one label is mentioned in a document, label all paragraphs as that label
-indicators refer to targets, and targets to goals, but this does not work the other way around
- check what happens with text in track changes
- Goals are to be found in texts about Millenium Development Goals. Should we include them?
- Still get Millennium goals if the word millennium is not present in the paragraph
- try to eliminate text with all the labels, since they're ambiguous
- should labels happening more than once be more weighted?
- E_CN.17_2007_14 -> goal 10 years later
- scraping 'https://sustainabledevelopment.un.org/sdg10'
- evaluate the very simple classifier that we have
'''





