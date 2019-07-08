import os
import PyPDF2
from bs4 import BeautifulSoup
import subprocess
from docx import Document as Docx
from utils.extract_utils import MAPPINGS, format_labels
from utils.text import Text
from langdetect import detect
import slate3k
import re



class Document:
    """
    Doc containing all paragraphs from .doc, .docx, .pdf
    """
    def __init__(self, filepath, filename):
        self.paragraphs = []
        self.name = filename
        self.filename_label = []
        self.from_any(filepath)
        self.doc_label_from(filename)

    def get_filename_label(self):
        return self.filename_label

    def doc_label_from(self, filename):
        # If 'goal|SDG no.' found in filename, process all doc with that label
        pattern =  r'(sdg|goal)\W?((?:[1-9]\b|1[0-7]?\b))'
        match = re.search(pattern, filename, flags=re.IGNORECASE)
        if match:
            self.filename_label = [int(match.group(2))]

    def from_any(self, filepath):
        base = os.path.basename(filepath)
        extension = os.path.splitext(base)[1].lower()
        if extension == '.doc':
            self.from_word(filepath)
        elif extension == '.pdf':
            self.from_pdf(filepath)
        elif extension == '.html':
            self.from_html(filepath)
        elif extension == '.txt':
            self.from_txt(filepath)
        else:
            print(f"File with extension {extension} not read.")

    def from_word(self, file):
        try:
            # Docx
            paragraphs = Docx(file).paragraphs
            for paragraphs in paragraphs:
                self.paragraphs.append(Text(paragraphs))
        except:
            try:
                # Doc
                string = subprocess.check_output(['antiword', '-t', file])
                # Decode and split by paragraphs
                extracted_list = string.decode('utf-8').split('\n\n')
                for paragraph in extracted_list:
                    self.paragraphs.append(Text(paragraph))
            except:
                # If Antiword does not work, convert to txt
                subprocess.run(['textutil', '-convert', 'txt', file])
                file = file.replace('.doc', '.txt')
                with open(file, 'r') as f:
                    data = f.readlines()
                    for paragraph in data:
                        self.paragraphs.append(Text(paragraph))

    def from_pdf(self, file):
        try:
            with open(file, 'rb') as fi:
                doc = slate3k.PDF(fi, word_margin=0)
                for i in range(len(doc)):
                    string = doc[i]
                    extracted_list = string.split('. \n')
                    for line in extracted_list:
                        self.paragraphs.append(Text(line))
        except:
            pass

    def from_html(self, file):
        with open(file, 'rb') as f:
            soup = BeautifulSoup(f, 'html.parser')
            # Strip out any code from the text
            for script in soup(["script", "style"]):
                script.decompose()
            for paragraph in soup.stripped_strings:
                try:
                    if detect(paragraph) == 'en':
                        self.paragraphs.append(Text(paragraph))
                except:
                    pass

    def from_txt(self, file):
        with open(file, 'r') as f:
            data = f.readlines()
            for paragraph in data:
                self.paragraphs.append(Text(paragraph))


def extract_labels(doc, q):
    labelled = {}
    unlabelled = {}
    patterns = [
                MAPPINGS['g'] + r"\W*\s+(,?\s*\b\d{1,2}\b[and\s\b\d{1,2}\b]*)",
                MAPPINGS['t'] + r"(\s+\d+\.[\d+a-d])",
                MAPPINGS['i'] + r"(\s+\d+\.[\d+a-d]\.\d+)"
            ]

    for paragraph in doc.paragraphs:
        labels = []
        goals = []
        text = paragraph.text

        # To avoid extracting Millennium Goals
        if 'millennium' in text.lower():
            pass
        elif ' mdg ' in text.lower():
            pass
        else:
            labelled_text = False
            for pattern in patterns:
                goals_extracted = re.findall(pattern, text, re.I)
                goals.extend(goals_extracted)
            text_labels = format_labels(goals)

            # Use labels from text if available
            if text_labels:
                if text not in labelled:
                    labelled[text] = {
                        'cats': text_labels,
                        'doc_id': doc.name
                        }
                else:
                    labelled[text]['cats'].extend(text_labels)
                labelled[text]['cats'] = list(set(labelled[text]['cats']))
                labelled_text = True

            # Use filename label is available and no labels found in text
            elif doc.filename_label:
                if text not in labelled:
                    labelled[text] = {
                            'cats': doc.filename_label,
                            'doc_id': doc.name
                            }
                else:
                    labelled[text]['cats'].extend(text_labels)
                labelled[text]['cats'] = list(set(labelled[text]['cats']))
                labelled_text = True

            if labelled_text == False:
                unlabelled[text] = None

    q.put((labelled, unlabelled))
