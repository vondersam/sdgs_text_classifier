import os
import PyPDF2
from bs4 import BeautifulSoup
import subprocess
import re
from docx import Document as Docx
from utils.extract_utils import MAPPINGS, format_labels



class Text:
    def __init__(self, t):
        self.text = t


class Document:
    """
    Doc containing all paragraphs from .doc, .docx, .pdf
    """
    def __init__(self, filepath):
        self.paragraphs = []
        self.from_any(filepath)


    def from_any(self, filepath):
        base = os.path.basename(filepath)
        extension = os.path.splitext(base)[1].lower()
        if extension == '.doc':
            self.from_word(filepath)
        elif extension == '.pdf':
            self.from_pdf(filepath)
        elif extension == '.html':
            self.from_html(filepath)
        else:
            print(f"File with extension {extension} not read.")

    def from_word(self, file):
        try:
            # Docx
            self.paragraphs = Docx(file).paragraphs
        except:
            try:
                # Doc
                extracted_string = subprocess.check_output(['antiword', '-t', file])
                # Decode and split by paragraphs
                extracted_list = extracted_string.decode('utf-8').split('\n\n')
                for paragraph in extracted_list:
                    # Eliminate extra spaces, replace space characters and eliminate break lines
                    processed_paragraph = re.sub(' +', ' ', paragraph.strip().replace('\xa0', ' ').replace('\n', ' '))
                    self.paragraphs.append(Text(processed_paragraph))
            except Exception as e:
                print(e)

    def from_pdf(self, file):
        try:
            with open(file, 'rb') as fi:
                pdfReader = PyPDF2.PdfFileReader(fi)
                for i in range(pdfReader.numPages):
                    extracted_string = pdfReader.getPage(i).extractText()
                    extracted_list = extracted_string.split('\n\n')
                    for line in extracted_list:
                        # Eliminate extra spaces, replace space characters and eliminate break lines
                        processed_text = re.sub(' +', ' ', line.strip().replace('\xa0', ' ').replace('\n', ' '))
                        self.paragraphs.append(Text(processed_text))
        except Exception as e:
            print(e)

    def from_html(self, file):
        with open(file, 'rb') as f:
            soup = BeautifulSoup(f, 'html.parser')
            for paragraph in soup.stripped_strings:
                self.paragraphs.append(Text(paragraph))

    def extract_labels(self):
        labelled_data = {}
        for paragraph in self.paragraphs:
            text = paragraph.text
            # To avoid extracting Millennium Goals
            if 'millennium' not in text.lower():
                patterns = [
                    MAPPINGS['g'] + r"\s+(,?\s*\b\d{1,2}\b[and\s\b\d{1,2}\b]*)",
                    MAPPINGS['t'] + r"(\s+\d+\.[\d+a-d])",
                    MAPPINGS['i'] + r"(\s+\d+\.[\d+a-d]\.\d+)"
                ]
                for pattern in patterns:
                    goals = re.findall(pattern, text, re.I)

                    for type_, numbers in goals:
                        labels = format_labels(type_, numbers)
                        if labels:
                            if text not in labelled_data:
                                labelled_data[text] = {
                                    'cats': labels
                                    #'labels': [],
                                    #'doc_id': document
                                }
                            labelled_data[text]['cats'].extend(labels)
                            labelled_data[text]['cats'] = list(set(labelled_data[text]['cats']))

        #print(labelled_data)
        return labelled_data
        #if document not in doc_tracker:
        #    doc_tracker[document] = []
        #doc_tracker[document].extend(labels)
